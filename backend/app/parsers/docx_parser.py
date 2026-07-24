"""DOCX text extraction (python-docx) — see app/parsers/registry.py for dispatch."""

import io

from docx import Document as DocxDocument

from app.parsers.base import ParsedDocument, ParsedSection


def parse(file_bytes: bytes) -> ParsedDocument:
    """Extract paragraph text and metadata from a DOCX's raw bytes.

    DOCX has no reliable page boundaries without rendering the layout (unlike
    PDF/PPTX), so this returns a single section rather than fabricating page
    numbers — `page_count` is left unset for the same reason.
    """
    document = DocxDocument(io.BytesIO(file_bytes))

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    title = document.core_properties.title or None
    sections = [ParsedSection(index=1, text=text)] if text else []

    return ParsedDocument(text=text, sections=sections, page_count=None, title=title)
